# ============================================================ 
# 📡 BASE DE DATOS DE MEDICIONES RNI - ENACOM
# Desarrollado por Lucas N. Miño y colaboradores
# ============================================================
# Este sistema permite cargar, procesar y visualizar mediciones
# de Radiaciones No Ionizantes (RNI) provenientes de archivos Excel
# estandarizados, generando resúmenes estadísticos, mapas, informes
# y exportaciones automáticas en formato Excel o Word.
# ============================================================

import streamlit as st
import pandas as pd
import numpy as np
from PIL import Image
from datetime import datetime, timedelta
import os, re, sqlite3   # >>> CAMBIO SQLITE (sacamos pickle, agregamos sqlite3)
from streamlit import rerun
import openpyxl
from openpyxl.styles import Alignment, Font
from openpyxl.drawing.image import Image as XLImage
import pydeck as pdk
import plotly.express as px
import plotly.graph_objects as go
import plotly.io as pio
from docx import Document
from docx.shared import Inches
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import A4
from io import BytesIO

# ---------------------- ESTILO ----------------------
with open("style.css") as f:
    st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)

# ---------------------- CONFIG ----------------------
st.set_page_config(page_title="Base de datos Radiaciones No Ionizantes - ENACOM v3.2", layout="wide")

# Logo institucional
try:
    logo = Image.open("logo_enacom.png")
    st.sidebar.image(logo, width=200)
except:
    st.sidebar.write("Logo ENACOM no encontrado")

# ------------------- SESSION STATE ------------------
st.session_state.setdefault("tabla_maestra", pd.DataFrame())
st.session_state.setdefault("uploaded_files_list", [])
st.session_state.setdefault("form_ccte", "")
st.session_state.setdefault("form_provincia", "")
st.session_state.setdefault("form_localidad", "")
st.session_state.setdefault("form_expediente", "")

DB_FILE = "rni.db"
TABLE_NAME = "tabla_maestra"

EXPECTED_COLS = [
    "CCTE", "Provincia", "Localidad",
    "Resultado", "Fecha", "Hora",
    "Nombre Archivo", "Expediente",
    "Sonda", "Lat", "Lon",
    "FechaCarga",
]
   
def load_tabla_maestra_from_db() -> pd.DataFrame:
    """Carga tabla_maestra desde SQLite. Si no existe, devuelve DF vacío."""
    if not os.path.exists(DB_FILE):
        return pd.DataFrame()
    conn = sqlite3.connect(DB_FILE)
    try:
        # Si la tabla no existe, devolvemos vacío
        tablas = pd.read_sql(
            "SELECT name FROM sqlite_master WHERE type='table';", conn
        )["name"].tolist()
        if TABLE_NAME not in tablas:
            # Si solo hay una tabla, la usamos igual
            if len(tablas) == 1:
                tabla_real = tablas[0]
            else:
                # No sabemos cuál es -> devolvemos vacío y dejamos que el usuario cargue de nuevo
                return pd.DataFrame()
        else:
            tabla_real = TABLE_NAME

        df = pd.read_sql(f"SELECT * FROM {tabla_real}", conn)

        # --- Normalizamos nombres de columnas (ccte / CCTE / CCTE_ / etc.) ---
        col_map = {}
        for c in df.columns:
            key = c.strip().lower().replace("ó", "o").replace("í", "i")
            if key == "ccte":
                col_map[c] = "CCTE"
            elif key == "provincia":
                col_map[c] = "Provincia"
            elif key == "localidad":
                col_map[c] = "Localidad"
            elif key in ("resultado", "resultado_con_incertidumbre"):
                col_map[c] = "Resultado"
            elif key == "fecha":
                col_map[c] = "Fecha"
            elif key in ("hora", "time"):
                col_map[c] = "Hora"
            elif key in ("nombrearchivo", "nombre_archivo", "archivo"):
                col_map[c] = "Nombre Archivo"
            elif key == "expediente":
                col_map[c] = "Expediente"
            elif key in ("sonda", "sonda_utilizada"):
                col_map[c] = "Sonda"
            elif key in ("lat", "latitud"):
                col_map[c] = "Lat"
            elif key in ("lon", "longitud"):
                col_map[c] = "Lon"
            elif key.lower() in ("fechacarga", "fecha_carga"):
                col_map[c] = "FechaCarga"

        if col_map:
            df = df.rename(columns=col_map)

        # Creamos columnas faltantes como NaN para que el resto del código no explote
        for col in EXPECTED_COLS:
            if col not in df.columns:
                df[col] = np.nan

        return df

    finally:
        conn.close()

def save_tabla_maestra_to_db(df: pd.DataFrame):
    """Guarda toda la tabla_maestra en SQLite, reemplazando el contenido."""
    if df is None:
        return
    conn = sqlite3.connect(DB_FILE)
    try:
        df.to_sql(TABLE_NAME, conn, if_exists="replace", index=False)
    finally:
        conn.close()

# Carga persistente de tabla maestra desde SQLite (ya no usamos PKL)
if st.session_state["tabla_maestra"].empty:
    try:
        st.session_state["tabla_maestra"] = load_tabla_maestra_from_db()
    except Exception as e:
        st.warning(f"No se pudo cargar tabla desde {DB_FILE}: {e}")

# ============================================================
# 🧩 FUNCIONES AUXILIARES
# ============================================================

def format_timedelta_long(td: timedelta) -> str:
    """Convierte un timedelta a formato hh:mm:ss."""
    total_seconds = int(td.total_seconds())
    hours = total_seconds // 3600
    minutes = (total_seconds % 3600) // 60
    seconds = total_seconds % 60
    return f"{hours:02d}:{minutes:02d}:{seconds:02d}"

def parse_dms_to_decimal(val):
    """Convierte coordenadas DMS (grados, minutos, segundos) a decimal."""
    if pd.isna(val):
        return np.nan
    try:
        return float(val)
    except:
        pass
    s = str(val).strip().replace(",", ".")
    m = re.search(r'([+-]?\d+(?:\.\d+)?)\D+(\d+(?:\.\d+)?)\D+(\d+(?:\.\d+)?)\D*\s*([NnSsEeWwOo])?', s)
    if m:
        d = float(m.group(1)); mnt = float(m.group(2)); sec = float(m.group(3))
        hemi = (m.group(4) or "").upper()
        dec = abs(d) + mnt/60.0 + sec/3600.0
        if hemi in ("S","W","O"):
            dec = -dec
        return dec
    m2 = re.search(r'([-+]?\d+(?:\.\d+)?(?:[eE][-+]?\d+)?)', s)
    if m2:
        try:
            return float(m2.group(1))
        except:
            return np.nan
    return np.nan

def extract_numeric_from_text(series):
    """Extrae valores numéricos (float) desde texto."""
    s = series.astype(str).str.replace(",", ".", regex=False)
    num = s.str.extract(r'([-+]?\d+(?:\.\d+)?(?:[eE][-+]?\d+)?)', expand=False)
    return pd.to_numeric(num, errors="coerce")

def find_index_column(df):
    """Detecta la columna que actúa como índice numérico."""
    candidates = ["índice", "indice", "index", "nro", "nº", "n°", "num", "numero", "#"]
    for c in df.columns:
        if any(cand in str(c).lower() for cand in candidates):
            return c
    return None

def calcular_tiempo_total_por_archivo(df: pd.DataFrame) -> timedelta:
    """
    Calcula la duración total de medición por archivo (considerando saltos de día),
    usando siempre Fecha + Hora y agrupando por 'Nombre Archivo'.
    """
    total = timedelta(0)
    if "Nombre Archivo" in df.columns:
        for _, df_archivo in df.groupby("Nombre Archivo"):
            if "Fecha" in df_archivo.columns and "Hora" in df_archivo.columns:
                for fecha, df_dia in df_archivo.groupby("Fecha"):
                    horas_validas = df_dia["Hora"].dropna()
                    if not horas_validas.empty:
                        dt_inicio = datetime.combine(fecha, horas_validas.min())
                        dt_fin = datetime.combine(fecha, horas_validas.max())
                        delta = dt_fin - dt_inicio
                        if delta.total_seconds() < 0:
                            delta += timedelta(days=1)
                        total += delta
    return total

# ============================================================
# ⚙️ PROCESAMIENTO DE ARCHIVOS EXCEL
# ============================================================

def procesar_archivos(uploaded_files, ccte, provincia, localidad, expediente):
    """Procesa múltiples archivos Excel y los integra en la tabla maestra."""
    lista_procesados, resumen_archivos = [], []

    for file in uploaded_files:
        try:
            df = pd.read_excel(file, header=8, engine="openpyxl")
        except Exception as e:
            st.warning(f"No se pudo leer {file.name}: {e}")
            continue

        df = df.dropna(axis=1, how="all")
        idx_col = find_index_column(df)
        total_mediciones = len(df)

        # Detecta número de mediciones
        if idx_col:
            df["_idx_num"] = pd.to_numeric(df[idx_col], errors="coerce")
            df = df[df["_idx_num"].notna()]
            if not df.empty:
                total_mediciones = int(df["_idx_num"].max())

        # Mapeo de columnas esperadas
        mapping_candidates = {
            "Fecha": ["fecha"],
            "Hora": ["hora", "time"],
            "Resultado": ["resultado con incertidumbre", "resultado"],
            "Sonda": ["sonda", "sonda utilizada"],
            "Lat": ["latitud", "lat"],
            "Lon": ["longitud", "lon"]
        }

        columnas_map, missing = {}, False
        for key, cands in mapping_candidates.items():
            found = next((c for c in df.columns if any(cand in str(c).lower() for cand in cands)), None)
            if not found and key not in ("Lat", "Lon"):
                st.warning(f"Archivo {file.name}: no se encontró columna para '{key}'")
                missing = True
                break
            if found:
                columnas_map[key] = found
        if missing:
            continue

        # Renombrado y limpieza
        df = df.rename(columns={v: k for k, v in columnas_map.items()})
        df["CCTE"], df["Provincia"], df["Localidad"] = ccte, provincia, localidad
        df["Expediente"] = expediente if expediente else os.path.splitext(file.name)[0]
        df["Nombre Archivo"] = file.name

        # Limpieza y formateo de campos
        if "Resultado" in df.columns:
            df["Resultado"] = extract_numeric_from_text(df["Resultado"])
        if "Lat" in df.columns:
            df["Lat"] = df["Lat"].apply(parse_dms_to_decimal)
        if "Lon" in df.columns:
            df["Lon"] = df["Lon"].apply(parse_dms_to_decimal)
        df.drop(columns=["_idx_num"], errors="ignore", inplace=True)

        lista_procesados.append(df)
        resumen_archivos.append({
            "archivo": file.name,
            "expediente": df["Expediente"].iloc[0],
            "total mediciones": total_mediciones,
            "max_resultado": df["Resultado"].max() if "Resultado" in df.columns else None
        })

    if lista_procesados:
        return pd.concat(lista_procesados, ignore_index=True), pd.DataFrame(resumen_archivos)
    return pd.DataFrame(), pd.DataFrame()

# ============================================================
# 🧹 FUNCIONES ADMINISTRATIVAS
# ============================================================

def eliminar_localidad(nombre_localidad: str):
    """Elimina una localidad completa de la tabla maestra."""
    if st.session_state["tabla_maestra"].empty:
        st.warning("⚠️ No hay datos cargados en la tabla maestra.")
        return

    df = st.session_state["tabla_maestra"]
    if "Localidad" not in df.columns:
        st.error("❌ No se encontró columna 'Localidad'.")
        return

    eliminados = len(df[df["Localidad"] == nombre_localidad])
    if eliminados == 0:
        st.info(f"ℹ️ No se encontró la localidad **{nombre_localidad}** en la tabla.")
        return

    st.session_state["tabla_maestra"] = df[df["Localidad"] != nombre_localidad]
    # >>> CAMBIO SQLITE: guardamos en DB
    save_tabla_maestra_to_db(st.session_state["tabla_maestra"])
    st.success(f"✅ Localidad **{nombre_localidad}** eliminada ({eliminados} registros).")

# ============================================================
# 📥 CARGA DE ARCHIVOS (SIDEBAR)
# ============================================================

st.sidebar.header("Cargar archivos")

if "uploader_key" not in st.session_state:
    st.session_state["uploader_key"] = 0

def reset_form():
    """Reinicia los campos del formulario lateral."""
    for key in ["uploaded_files_list", "form_localidad", "form_expediente", "form_ccte", "form_provincia"]:
        st.session_state[key] = "" if "list" not in key else []
    st.session_state["uploader_key"] += 1
    rerun()

# --- Formulario lateral de carga ---
with st.sidebar.form("carga_form", clear_on_submit=False):
    ccte = st.selectbox("CCTE", ["CABA", "Buenos Aires", "Comodoro Rivadavia", "Córdoba", "Neuquén", "Posadas", "Salta"], key="form_ccte")
    provincia = st.selectbox(
        "Provincia",
        ["Buenos Aires","CABA","Catamarca","Chaco","Chubut","Córdoba","Corrientes","Entre Ríos","Formosa","Jujuy",
         "La Pampa","La Rioja","Mendoza","Misiones","Neuquén","Río Negro","Salta","San Juan","San Luis","Santa Cruz",
         "Santa Fe","Santiago del Estero","Tierra del Fuego","Tucumán"],
        key="form_provincia"
    )
    localidad = st.text_input("Localidad", value=st.session_state["form_localidad"], key="form_localidad")
    expediente = st.text_input("Expediente", value=st.session_state["form_expediente"], key="form_expediente")
    files = st.file_uploader("Seleccionar archivos Excel", accept_multiple_files=True, type=["xlsx"], key=f"form_files_{st.session_state['uploader_key']}")
    submit = st.form_submit_button("Procesar archivos")

    if submit and files:
        df_proc, resumen_df = procesar_archivos(files, ccte, provincia, localidad, expediente)
        if not df_proc.empty:
            df_proc["FechaCarga"] = datetime.now()
            st.session_state["tabla_maestra"] = pd.concat([st.session_state["tabla_maestra"], df_proc], ignore_index=True)
            # >>> CAMBIO SQLITE: guardamos en DB
            save_tabla_maestra_to_db(st.session_state["tabla_maestra"])
            st.success(f"{len(files)} archivos procesados y agregados.")
            st.sidebar.dataframe(resumen_df)
            st.session_state["uploader_key"] += 1
        else:
            st.warning("No se procesaron archivos válidos.")

st.sidebar.button("Restablecer formulario", on_click=reset_form)

# ------------------- SIDEBAR: eliminar localidad ------------------
if "tabla_maestra" in st.session_state and not st.session_state["tabla_maestra"].empty:
    localidades_unicas = sorted(st.session_state["tabla_maestra"]["Localidad"].dropna().unique().tolist())
    localidad_a_borrar = st.sidebar.selectbox("Seleccionar localidad a eliminar", [""] + localidades_unicas)

    if st.sidebar.button("❌ Eliminar localidad") and localidad_a_borrar:
        eliminar_localidad(localidad_a_borrar)

# ------------------- ENCABEZADO CON LOGO ------------------
col1, col2 = st.columns([6,1])

with col1:
    st.title(" ")

with col2:
    st.image("logo_enacom.png")

# ------------------- HIGHLIGHT GLOBAL ------------------
if "tabla_maestra" in st.session_state and not st.session_state["tabla_maestra"].empty:
    df = st.session_state["tabla_maestra"].copy()
    df["Resultado"] = pd.to_numeric(df["Resultado"], errors="coerce")
    idx_max = df["Resultado"].idxmax()
    fila_max = df.loc[idx_max]

    localidad_top = fila_max.get("Localidad", "N/A")
    resultado_top = fila_max["Resultado"]
    resultado_top_pct = resultado_top**2 / 3770 / 0.20021 * 100 if pd.notna(resultado_top) else None

    # Fecha/Hora asociada
    fecha_top = None
    if "FechaHora" in fila_max and pd.notna(fila_max["FechaHora"]):
        fecha_top = fila_max["FechaHora"]
    elif "Fecha" in fila_max and "Hora" in fila_max:
        try:
            fecha_top = datetime.combine(fila_max["Fecha"], fila_max["Hora"])
        except:
            fecha_top = fila_max["Fecha"]
    elif "Fecha" in fila_max:
        fecha_top = fila_max["Fecha"]

    st.markdown("## 🌎 Valor máximo registrado en Argentina")
    # --- Estilo visual con CSS ---
    st.markdown("""
    <style>
    div[data-testid="stMetricContainer"] {
        background: rgba(240, 248, 255, 0.6);
        border: 1px solid rgba(200, 200, 200, 0.3);
        border-radius: 12px;
        padding: 16px;
        text-align: center;
        box-shadow: 0 1px 6px rgba(0,0,0,0.1);
        transition: all 0.2s ease-in-out;
    }
    div[data-testid="stMetricContainer"]:hover {
        transform: translateY(-2px);
        box-shadow: 0 4px 10px rgba(0,0,0,0.15);
    }
    div[data-testid="stMetricLabel"] > div {
        font-size: 16px;
        font-weight: 600;
        color: #2E3B55;
    }
    div[data-testid="stMetricValue"] {
        font-size: 26px;
        font-weight: 700;
        color: #004aad;
    }
    </style>
    """, unsafe_allow_html=True)

    col1, col2, col3, col4 = st.columns([2.5, 1.5, 1.5, 1.5])
    col1.metric("Localidad", localidad_top)
    col2.metric("Resultado máximo V/m", f"{resultado_top:.2f}")
    col3.metric("Resultado máximo (%)", f"{resultado_top_pct:.2f}" if resultado_top_pct else "N/A")
    col4.metric("Fecha/Hora", str(fecha_top))

# ------------------- RESUMEN GENERAL DE LOCALIDADES (con filtros previos) ------------------
if "tabla_maestra" in st.session_state and not st.session_state["tabla_maestra"].empty:
    df = st.session_state["tabla_maestra"].copy()
    df["Resultado"] = pd.to_numeric(df["Resultado"], errors="coerce")

    # --- 🔍 FILTROS PREVIOS ---    
    st.header("📊 Resumen general de mediciones")

    col1, col2, col3 = st.columns([1, 1, 1])
    with col1:
        ccte_sel = st.selectbox(
            "Filtrar CCTE",
            ["Todos"] + sorted(df["CCTE"].dropna().unique().tolist()),
            key="resumen_ccte"
        )
        if ccte_sel != "Todos":
            df = df[df["CCTE"] == ccte_sel]

    with col2:
        prov_sel = st.selectbox(
            "Filtrar Provincia",
            ["Todas"] + sorted(df["Provincia"].dropna().unique().tolist()),
            key="resumen_provincia"
        )
        if prov_sel != "Todas":
            df = df[df["Provincia"] == prov_sel]

    with col3:
        año_sel = "Todos"
        if "Fecha" in df.columns:
            df["Fecha"] = pd.to_datetime(df["Fecha"], dayfirst=True, errors="coerce")
            años_disp = sorted(df["Fecha"].dt.year.dropna().astype(int).unique().tolist(), reverse=True)
            if años_disp:
                año_sel = st.selectbox(
                    "Filtrar Año",
                    ["Todos"] + [str(a) for a in años_disp],
                    key="resumen_año"
                )
                if año_sel != "Todos":
                    df = df[df["Fecha"].dt.year == int(año_sel)]

    # --- Procesamiento base ---
    if "Fecha" in df.columns:
        df["Fecha"] = pd.to_datetime(df["Fecha"], dayfirst=True, errors='coerce').dt.date
    if "Hora" in df.columns:
        df["Hora"] = pd.to_datetime(df["Hora"], errors='coerce').dt.time

    if "Fecha" in df.columns and "Hora" in df.columns:
        df["FechaHora"] = df.apply(
            lambda x: datetime.combine(x["Fecha"], x["Hora"]) if pd.notna(x["Fecha"]) and pd.notna(x["Hora"]) else pd.NaT,
            axis=1
        )
    else:
        df["FechaHora"] = pd.NaT

    # --- Resumen agrupado ---
    resumen_localidad = []
    for (ccte, prov, loc), g in df.groupby(["CCTE", "Provincia", "Localidad"]):
        inicio = g["FechaHora"].min() if "FechaHora" in g.columns else None
        fin = g["FechaHora"].max() if "FechaHora" in g.columns else None
        tiempo_total_localidad = calcular_tiempo_total_por_archivo(g)
        max_res = g["Resultado"].max() if pd.notna(g["Resultado"].max()) else None
        resumen_localidad.append({
            "CCTE": ccte,
            "Provincia": prov,
            "Localidad": loc,
            "Inicio": inicio,
            "Fin": fin,
            "Mediciones": len(g),
            "Tiempo mediciones": format_timedelta_long(tiempo_total_localidad),
            "Resultado Max (V/m)": max_res,
            "Resultado Max (%)": max_res**2 / 3770 / 0.20021 * 100 if max_res else None,
            "N° Expediente": ", ".join(sorted(g["Expediente"].dropna().unique().astype(str))),
            "Sonda utilizada": ", ".join(sorted(g["Sonda"].dropna().unique().astype(str))) if "Sonda" in g.columns else "N/A"
        })

    resumen_localidad_df = pd.DataFrame(resumen_localidad)
    
    st.dataframe(resumen_localidad_df)

    # --- Botón de exportación ---
    if not resumen_localidad_df.empty:
        if st.button("📥 Exportar resumen filtrado a Excel"):
            try:
                ruta_excel = "resumen_localidades_filtrado.xlsx"
                resumen_localidad_df.to_excel(ruta_excel, index=False)
                wb = openpyxl.load_workbook(ruta_excel)
                ws = wb.active

                # Logo institucional
                try:
                    logo_path = "logo_enacom.png"
                    img = XLImage(logo_path)
                    img.width, img.height = 200, 70
                    ws.add_image(img, "A1")
                    ws.insert_rows(1, amount=5)
                except Exception as e:
                    st.warning(f"No se pudo insertar logo: {e}")

                for cell in ws[6]:
                    cell.font = Font(bold=True)
                    cell.alignment = Alignment(horizontal="center", vertical="center")

                for row in ws.iter_rows(min_row=7, max_row=ws.max_row, min_col=1, max_col=ws.max_column):
                    for cell in row:
                        cell.alignment = Alignment(horizontal="center", vertical="center")

                wb.save(ruta_excel)
                st.success(f"Archivo '{ruta_excel}' generado con formato y logo.")
                with open(ruta_excel, "rb") as f:
                    st.download_button(
                        label="⬇️ Descargar Excel filtrado",
                        data=f,
                        file_name=ruta_excel,
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )
            except Exception as e:
                st.error(f"Error exportando Excel: {e}")

#----------------------------- GRAFICOS-------------------------------------
if not st.session_state["tabla_maestra"].empty:
    df_grafico = st.session_state["tabla_maestra"].copy()
    
    # Distribución de puntos medidos por CCTE
    df_pie = df_grafico.groupby("CCTE").size().reset_index(name="Cantidad Puntos")
    fig_pie = px.pie(
        df_pie,
        names="CCTE",
        values="Cantidad Puntos",
        title="Distribución de puntos medidos por CCTE",
        color_discrete_sequence=px.colors.qualitative.Set3
    )

    # Localidades por Provincia y CCTE
    resumen = df_grafico.groupby(["Provincia","CCTE"])["Localidad"].nunique().reset_index(name="CantidadLocalidades")
    fig_bar = px.bar(
        resumen,
        x="Provincia",
        y="CantidadLocalidades",
        color="CCTE",
        text="CantidadLocalidades",
        barmode="group",
        title="Localidades por Provincia y CCTE"
    )

    st.subheader("📊 Resumen de mediciones y localidades")
    col1, col2 = st.columns(2)
    with col1:
        st.plotly_chart(fig_pie, width="stretch")
    with col2:
        st.plotly_chart(fig_bar, width="stretch")

# ------------------- RESUMEN Y EDICIÓN DE LOCALIDAD ------------------
st.header("📊 Gestión de Localidades")

df_base = st.session_state["tabla_maestra"].copy()

columnas_necesarias = {"CCTE", "Provincia", "Localidad"}
if df_base.empty or not columnas_necesarias.issubset(df_base.columns):
    st.info("Todavía no hay datos suficientes (o faltan columnas CCTE/Provincia/Localidad) para gestionar localidades. Cargá mediciones nuevas.")
    df_filtrado_prov = pd.DataFrame()
    localidad_seleccionada = ""
    provincia_filtro = "Todas"
    ccte_filtro = "Todos"
else:
    col1, col2, col3, col4 = st.columns([1, 1, 1, 1])

    with col1:
        lista_ccte = sorted(df_base["CCTE"].dropna().unique().tolist())
        ccte_filtro = st.selectbox(
            "Filtrar CCTE",
            ["Todos"] + lista_ccte,
            key="gestion_ccte"
        )
        if ccte_filtro == "Todos":
            df_filtrado_ccte = df_base.copy()
        else:
            df_filtrado_ccte = df_base[df_base["CCTE"] == ccte_filtro].copy()

    with col2:
        lista_prov = sorted(df_filtrado_ccte["Provincia"].dropna().unique().tolist())
        provincia_filtro = st.selectbox(
            "Filtrar Provincia",
            ["Todas"] + lista_prov,
            key="gestion_provincia"
        )
        if provincia_filtro == "Todas":
            df_filtrado_prov = df_filtrado_ccte.copy()
        else:
            df_filtrado_prov = df_filtrado_ccte[df_filtrado_ccte["Provincia"] == provincia_filtro].copy()

    with col4:
        año_filtro = "Todos"
        if not df_filtrado_prov.empty and "Fecha" in df_filtrado_prov.columns:
            _años = pd.to_datetime(df_filtrado_prov["Fecha"], dayfirst=True, errors="coerce").dt.year
            df_filtrado_prov["_Año"] = _años
            años_disponibles = sorted(df_filtrado_prov["_Año"].dropna().astype(int).unique().tolist(), reverse=True)
            opciones_año = ["Todos"] + [str(a) for a in años_disponibles]
            año_filtro = st.selectbox("📅 Año", opciones_año, index=0, key="gestion_año")
            if año_filtro != "Todos":
                df_filtrado_prov = df_filtrado_prov[df_filtrado_prov["_Año"] == int(año_filtro)]
            df_filtrado_prov = df_filtrado_prov.drop(columns=["_Año"])

    with col3:
        if not df_filtrado_prov.empty:
            localidades_cargadas = df_filtrado_prov["Localidad"].dropna().unique().tolist()
        else:
            localidades_cargadas = []
        localidad_seleccionada = st.selectbox(
            "Seleccionar Localidad",
            [""] + sorted(localidades_cargadas),
            key="gestion_localidad"
        )

# Subset final
if localidad_seleccionada:
    df_localidad = df_filtrado_prov[df_filtrado_prov["Localidad"] == localidad_seleccionada].copy()
else:
    df_localidad = df_filtrado_prov.copy()

# Convertir fechas y horas
if "Fecha" in df_localidad.columns:
    df_localidad["Fecha"] = pd.to_datetime(df_localidad["Fecha"], dayfirst=True, errors='coerce').dt.date
if "Hora" in df_localidad.columns:
    df_localidad["Hora"] = pd.to_datetime(df_localidad["Hora"], errors='coerce').dt.time
if "Fecha" in df_localidad.columns and "Hora" in df_localidad.columns:
    df_localidad["FechaHora"] = df_localidad.apply(
        lambda x: datetime.combine(x["Fecha"], x["Hora"]) if pd.notna(x["Fecha"]) and pd.notna(x["Hora"]) else pd.NaT,
        axis=1
    )
else:
    df_localidad["FechaHora"] = pd.NaT    

# ---------------- Datos generales ----------------
if localidad_seleccionada:    
    provincia_real = df_localidad["Provincia"].iloc[0] if "Provincia" in df_localidad.columns else "N/A"
    titulo_scope = f"la localidad {localidad_seleccionada}, {provincia_real}"
elif provincia_filtro != "Todas":
    titulo_scope = f"{provincia_filtro}"
elif ccte_filtro != "Todos":
    titulo_scope = f"CCTE {ccte_filtro}"
else:
    titulo_scope = "todo el país"

st.subheader(f"Mediciones RNI de {titulo_scope}")

tiempo_total_localidad = calcular_tiempo_total_por_archivo(df_localidad)
total_puntos = len(df_localidad)
max_resultado = df_localidad["Resultado"].max() if "Resultado" in df_localidad.columns else None
max_resultado_pct = max_resultado**2 / 3770 / 0.20021 * 100 if pd.notna(max_resultado) else None
sondas = df_localidad["Sonda"].dropna().unique().tolist() if "Sonda" in df_localidad.columns else []

st.write(f"Cantidad total de puntos medidos: {total_puntos}")
st.write(f"Máximo Resultado (V/m): {max_resultado}")
st.write(f"Sonda utilizada: {', '.join(sondas) if sondas else 'N/A'}")
st.write(f"Tiempo total de mediciones: {format_timedelta_long(tiempo_total_localidad)} horas")

# ---------------- Resumen por día y mes ----------------
if "FechaHora" in df_localidad.columns and not df_localidad.empty:
    df_localidad["FechaHora"] = pd.to_datetime(df_localidad["FechaHora"])
    df_localidad["Fecha"] = df_localidad["FechaHora"].dt.date
    df_localidad["Mes"] = df_localidad["FechaHora"].dt.to_period("M")

    # --- Resumen diario ---
    def resumen_por_dia(df_dia):
        tiempo_total = calcular_tiempo_total_por_archivo(df_dia)
        inicio_dt = df_dia["FechaHora"].min()
        fin_dt = df_dia["FechaHora"].max()
        hora_inicio = inicio_dt.strftime("%H:%M:%S") if pd.notna(inicio_dt) else "-"
        hora_fin = fin_dt.strftime("%H:%M:%S") if pd.notna(fin_dt) else "-"
        puntos = len(df_dia)
        localidades = ", ".join(sorted(df_dia["Localidad"].dropna().unique()))
        return {
            "Hora de inicio": hora_inicio,
            "Hora de fin": hora_fin,
            "Tiempo total trabajado": format_timedelta_long(tiempo_total),
            "Cantidad de puntos medidos": puntos,
            "Localidades trabajadas (por día)": localidades,
        }

    filas_resumen_dias = []
    for fecha, g_dia in df_localidad.groupby("Fecha"):
        info = resumen_por_dia(g_dia)
        info["Fecha de medición"] = fecha
        filas_resumen_dias.append(info)

    resumen_dias = pd.DataFrame(filas_resumen_dias)
    if not resumen_dias.empty:
        resumen_dias = resumen_dias[
            [
                "Fecha de medición",
                "Hora de inicio",
                "Hora de fin",
                "Tiempo total trabajado",
                "Cantidad de puntos medidos",
                "Localidades trabajadas (por día)",
            ]
        ]
    
    # --- Resumen mensual ---
    resumen_mensual = df_localidad.groupby("Mes").agg({
        "FechaHora": ["min","max"],
        "Localidad": lambda x: ", ".join(sorted(x.dropna().unique())),
        "Resultado": "count"
    }).reset_index()

    resumen_mensual.columns = ["Mes","Hora inicio","Hora fin","Localidades trabajadas","Cantidad puntos"]

    # Calcular tiempo total trabajado por mes
    def calcular_tiempo_mes(g_mes):
        return format_timedelta_long(calcular_tiempo_total_por_archivo(g_mes))

    filas_tiempo_mes = []
    for mes, g_mes in df_localidad.groupby("Mes"):
        filas_tiempo_mes.append({
            "Mes": mes,
            "Horas trabajadas": calcular_tiempo_mes(g_mes)
        })
    tiempo_por_mes = pd.DataFrame(filas_tiempo_mes)

    resumen_mensual = resumen_mensual.merge(tiempo_por_mes, on="Mes")

    # -------- Tabs para elegir vista --------
    tab1, tab2, tab3 = st.tabs(["📅 Resumen Diario", "🗓️ Resumen Mensual", "📊 Gráfico"])

    with tab1:
        st.markdown(f"### ⏱️ Tiempo trabajado por día en {titulo_scope}")
        st.dataframe(resumen_dias)

    with tab2:
        st.markdown(f"### 📅 Mediciones Totales por mes en {titulo_scope}")
        st.dataframe(resumen_mensual)

    with tab3:
        if not resumen_mensual.empty:
            st.markdown(f"### 📊 Gráfico mensual de mediciones y tiempo trabajado en {titulo_scope}")

            # Crear figura con dos ejes: cantidad de puntos y horas trabajadas
            fig = go.Figure()

            # Barra: Cantidad de puntos
            fig.add_trace(go.Bar(
                x=resumen_mensual["Mes"].astype(str),
                y=resumen_mensual["Cantidad puntos"],
                name="Cantidad puntos",
                marker_color="steelblue",
                yaxis="y1",
                text=resumen_mensual["Cantidad puntos"],
                textposition="auto",
                hovertext=resumen_mensual["Localidades trabajadas"],  # 👈 tooltip
                hovertemplate="<b>%{x}</b><br>Puntos: %{y}<br>Localidades: %{hovertext}"
            ))

            # Línea: Horas trabajadas   
            def tiempo_a_horas(s):
                h, m, sec = map(int, s.split(":"))
                return h + m/60 + sec/3600

            resumen_mensual["Horas trabajadas num"] = resumen_mensual["Horas trabajadas"].apply(tiempo_a_horas)

            fig.add_trace(go.Scatter(
                x=resumen_mensual["Mes"].astype(str),
                y=resumen_mensual["Horas trabajadas num"],
                name="Horas trabajadas",
                yaxis="y2",
                mode="lines+markers",
                line=dict(color="orange", width=2)
            ))

            # Configuración de ejes
            fig.update_layout(
                xaxis=dict(title="Mes"),
                yaxis=dict(title="Cantidad de puntos", side="left"),
                yaxis2=dict(
                    title="Horas trabajadas",
                    overlaying="y",
                    side="right"
                ),
                legend=dict(x=0.01, y=0.99),
                template="plotly_white",
                height=450
            )

            st.plotly_chart(fig, width="stretch")

# ---------------- Semáforo ----------------
if max_resultado_pct and not df_localidad.empty:
    rangos_colores = [
        (0, 1, "#84C2F5"), (1, 2, "#489DFF"), (2, 4, "#006BD6"),
        (4, 8, "#A9E7A9"), (8, 15, "#89DD89"), (15, 20, "#4D9623"),
        (20, 35, "#D9FF00"), (35, 50, "#F39A6D"), (50, 100, "#E68200"),
        (100, float("inf"), "#CC0000")
    ]
    color_localidad = next((color for low, high, color in rangos_colores if low <= max_resultado_pct < high), "#FFFFFF")
    st.markdown(
        f"""
        <div style="
            background-color:{color_localidad};
            padding:20px;
            border-radius:10px;
            text-align:center;
            font-size:24px;
            font-weight:bold;
            color:#000;">
            Resultado máximo en {titulo_scope}: {max_resultado_pct:.2f} %
        </div>
        """,
        unsafe_allow_html=True
    )
    # Imagen del semáforo de colores 
    st.image("mapa de calor.png", caption="Escala de colores para interpretar los resultados", width="stretch")

# ------------------- MAPA INTERACTIVO ------------------
if "Lat" in df_localidad.columns and "Lon" in df_localidad.columns and not df_localidad.empty:
    coords = df_localidad.dropna(subset=["Lat", "Lon"])[["Lat", "Lon", "Localidad", "Resultado"]].copy()
    if not coords.empty:
        # Forzamos coordenadas negativas (Argentina)
        coords["lat"] = coords["Lat"].apply(lambda x: -abs(x))
        coords["lon"] = coords["Lon"].apply(lambda x: -abs(x))

        rangos_colores_map = [
            (0, 1, [132, 194, 245]),
            (1, 2, [72, 157, 255]),
            (2, 4, [0, 107, 214]),
            (4, 8, [169, 231, 169]),
            (8, 15, [137, 221, 137]),
            (15, 20, [77, 150, 35]),
            (20, 35, [217, 255, 0]),
            (35, 50, [243, 154, 109]),
            (50, 100, [230, 130, 0]),
            (100, float("inf"), [204, 0, 0])
        ]

        def color_semaforo(valor):
            if pd.isna(valor):
                return [200, 200, 200]
            for low, high, color in rangos_colores_map:
                if low <= valor < high:
                    return color
            return [0, 0, 0]

        coords["color"] = coords["Resultado"].apply(color_semaforo)

        st.subheader("🗺️ Mapa Semaforizado")
        mapa = pdk.Deck(
            map_style="https://basemaps.cartocdn.com/gl/positron-gl-style/style.json",
            initial_view_state=pdk.ViewState(
                latitude=coords["lat"].mean(),
                longitude=coords["lon"].mean(),
                zoom=6,
                pitch=0,
            ),
            layers=[
                pdk.Layer(
                    "ScatterplotLayer",
                    data=coords,
                    get_position='[lon, lat]',
                    get_fill_color='color',
                    get_radius=12,
                    pickable=True,
                )
            ],
            tooltip={"text": "Localidad: {Localidad}\nResultado: {Resultado}"}
        )
        st.pydeck_chart(mapa, width="stretch")

# -------------------- Edición de información (plegable) --------------------
if localidad_seleccionada:
    # >>> CAMBIO SQLITE: aseguramos que FechaCarga sea datetime antes de usar strftime
    if "FechaCarga" in st.session_state["tabla_maestra"].columns:
        st.session_state["tabla_maestra"]["FechaCarga"] = pd.to_datetime(
            st.session_state["tabla_maestra"]["FechaCarga"], errors="coerce"
        )

    ultima_fecha = None
    if "FechaCarga" in st.session_state["tabla_maestra"].columns:
        mask_fecha = st.session_state["tabla_maestra"]["Localidad"] == localidad_seleccionada
        if mask_fecha.any():
            ultima_fecha = st.session_state["tabla_maestra"].loc[mask_fecha, "FechaCarga"].max()

    expander_title = f"✏️ Editar información de {localidad_seleccionada}"
    if ultima_fecha is not None and pd.notna(ultima_fecha):
        expander_title += f" (Última modificación: {ultima_fecha.strftime('%d/%m/%Y %H:%M:%S')})"

    with st.expander(expander_title, expanded=False):
        ccte_actual = df_localidad["CCTE"].iloc[0]
        provincia_actual = df_localidad["Provincia"].iloc[0]
        localidad_actual = df_localidad["Localidad"].iloc[0]
        expediente_actual = df_localidad["Expediente"].iloc[0]

        nuevo_ccte = st.selectbox(
            "CCTE",
            ["CABA", "Buenos Aires", "Comodoro Rivadavia", "Córdoba", "Neuquén", "Posadas", "Salta"],
            index=["CABA","Buenos Aires","Comodoro Rivadavia","Córdoba","Neuquén","Posadas","Salta"].index(ccte_actual)
        )
        nueva_provincia = st.selectbox(
            "Provincia",
            ["Buenos Aires","CABA","Catamarca","Chaco","Chubut","Córdoba","Corrientes","Entre Ríos","Formosa","Jujuy",
             "La Pampa","La Rioja","Mendoza","Misiones","Neuquén","Río Negro","Salta","San Juan","San Luis","Santa Cruz",
             "Santa Fe","Santiago del Estero","Tierra del Fuego","Tucumán"],
            index=["Buenos Aires","CABA","Catamarca","Chaco","Chubut","Córdoba","Corrientes","Entre Ríos","Formosa","Jujuy",
                   "La Pampa","La Rioja","Mendoza","Misiones","Neuquén","Río Negro","Salta","San Juan","San Luis","Santa Cruz",
                   "Santa Fe","Santiago del Estero","Tierra del Fuego","Tucumán"].index(provincia_actual)
        )
        nueva_localidad = st.text_input("Localidad", value=localidad_actual)
        nuevo_expediente = st.text_input("Expediente", value=expediente_actual)

        if "FechaCarga" not in st.session_state["tabla_maestra"].columns:
            st.session_state["tabla_maestra"]["FechaCarga"] = pd.NaT

        def guardar_cambios():
            mask = st.session_state["tabla_maestra"]["Localidad"] == localidad_actual
            st.session_state["tabla_maestra"].loc[mask, "CCTE"] = nuevo_ccte
            st.session_state["tabla_maestra"].loc[mask, "Provincia"] = nueva_provincia
            st.session_state["tabla_maestra"].loc[mask, "Localidad"] = nueva_localidad
            st.session_state["tabla_maestra"].loc[mask, "Expediente"] = nuevo_expediente
            st.session_state["tabla_maestra"].loc[mask, "FechaCarga"] = datetime.now()

            try:
                # >>> CAMBIO SQLITE: guardamos en DB
                save_tabla_maestra_to_db(st.session_state["tabla_maestra"])
                st.success("Cambios guardados correctamente")
            except Exception as e:
                st.error(f"No se pudieron guardar los cambios: {e}")

        st.button("💾 Guardar cambios", on_click=guardar_cambios)

        def eliminar_localidad_cb():
            mask = st.session_state["tabla_maestra"]["Localidad"] == localidad_actual
            if mask.any():
                st.session_state["tabla_maestra"] = st.session_state["tabla_maestra"].loc[~mask]
                try:
                    # >>> CAMBIO SQLITE: guardamos en DB
                    save_tabla_maestra_to_db(st.session_state["tabla_maestra"])
                    st.success(f"Localidad '{localidad_actual}' eliminada correctamente")
                    st.experimental_rerun()  # recarga la app para reflejar cambios
                except Exception as e:
                    st.error(f"No se pudo eliminar la localidad: {e}")
            else:
                st.warning("No se encontró la localidad para eliminar.")

        st.button("🗑️ Eliminar localidad", on_click=eliminar_localidad_cb)

# ============================================================
# 🖨️ EXPORTACIÓN DE INFORMES PDF / WORD
# ============================================================

with st.expander("🖨️ Generar informe PDF / Word", expanded=False):
    st.header("🖨️ Generar Informe con Gráficos y Datos Resumidos")

    if not st.session_state["tabla_maestra"].empty:
        # Usamos el mismo subset que se está viendo en pantalla:
        df_export = df_localidad.copy() if not df_localidad.empty else df_filtrado_prov.copy()

        # Si por algún motivo ese df está vacío, fallback a tabla completa
        if df_export.empty:
            df_export = st.session_state["tabla_maestra"].copy()

        # ========= ESTADÍSTICAS PARA EL RELATO =========
        df_export["Resultado"] = pd.to_numeric(df_export.get("Resultado", np.nan), errors="coerce")

        max_resultado = df_export["Resultado"].max() if "Resultado" in df_export.columns else None
        max_resultado_pct = None
        localidad_max = provincia_max = ccte_max = "N/D"
        fecha_hora_max = None

        if pd.notna(max_resultado):
            max_resultado_pct = max_resultado**2 / 3770 / 0.20021 * 100

            fila_max = df_export.loc[df_export["Resultado"].idxmax()]
            localidad_max = fila_max.get("Localidad", "N/D")
            provincia_max = fila_max.get("Provincia", "N/D")
            ccte_max = fila_max.get("CCTE", "N/D")

            # Fecha y hora del máximo
            if "FechaHora" in fila_max and pd.notna(fila_max["FechaHora"]):
                fecha_hora_max = fila_max["FechaHora"]
            elif "Fecha" in fila_max and "Hora" in fila_max:
                try:
                    fecha_hora_max = datetime.combine(fila_max["Fecha"], fila_max["Hora"])
                except Exception:
                    fecha_hora_max = fila_max.get("Fecha", None)
            else:
                fecha_hora_max = fila_max.get("Fecha", None)

        # Rango de fechas trabajadas
        fecha_min = fecha_max_med = None
        if "Fecha" in df_export.columns:
            fechas = pd.to_datetime(df_export["Fecha"], dayfirst=True, errors="coerce")
            if fechas.notna().any():
                fecha_min = fechas.min().date()
                fecha_max_med = fechas.max().date()

        # Tiempo total trabajado (según Nombre Archivo + Fecha/Hora)
        tiempo_total_trabajado = calcular_tiempo_total_por_archivo(df_export)

        # Sondas utilizadas
        sondas_uniq = []
        if "Sonda" in df_export.columns:
            sondas_uniq = sorted(df_export["Sonda"].dropna().astype(str).unique().tolist())

        # ========= DESGLOSE POR MES =========
        resumen_mensual_export = pd.DataFrame()
        if "FechaHora" in df_export.columns and df_export["FechaHora"].notna().any():
            df_tmp = df_export.copy()
            df_tmp["FechaHora"] = pd.to_datetime(df_tmp["FechaHora"], errors="coerce")
            df_tmp["Mes"] = df_tmp["FechaHora"].dt.to_period("M")

            if df_tmp["Mes"].notna().any():
                resumen_mensual_export = df_tmp.groupby("Mes").agg(
                    Cantidad_puntos=("Resultado", "count"),
                    Fecha_inicio=("FechaHora", "min"),
                    Fecha_fin=("FechaHora", "max"),
                    Localidades_trabajadas=("Localidad", lambda x: ", ".join(sorted(x.dropna().unique())))
                ).reset_index()

                # Tiempo por mes
                def _tiempo_mes(g_mes):
                    return format_timedelta_long(calcular_tiempo_total_por_archivo(g_mes))

                filas_tiempo_export = []
                for mes, g_mes in df_tmp.groupby("Mes"):
                    filas_tiempo_export.append({
                        "Mes": mes,
                        "Horas_trabajadas": _tiempo_mes(g_mes)
                    })
                tiempos_mes = pd.DataFrame(filas_tiempo_export)

                resumen_mensual_export = resumen_mensual_export.merge(tiempos_mes, on="Mes")

        # ========= TABLA DE EXPEDIENTES =========
        expedientes_df = pd.DataFrame()
        if "Expediente" in df_export.columns:
            expedientes_df = df_export.groupby("Expediente").agg(
                Cantidad_puntos=("Resultado", "count"),
                CCTE=("CCTE", lambda x: ", ".join(sorted(x.dropna().unique()))),
                Provincias=("Provincia", lambda x: ", ".join(sorted(x.dropna().unique()))),
                Localidades=("Localidad", lambda x: ", ".join(sorted(x.dropna().unique()))),
                Max_Vm=("Resultado", "max")
            ).reset_index()
            expedientes_df = expedientes_df.sort_values(by="Max_Vm", ascending=False)

        # ========= GRÁFICO SOLO DEL ÁMBITO ACTUAL =========
        df_graf_export = df_export.copy()
        if {"Provincia", "CCTE", "Localidad"}.issubset(df_graf_export.columns):
            resumen_export = (
                df_graf_export
                .groupby(["Provincia", "CCTE"])["Localidad"]
                .nunique()
                .reset_index(name="CantidadLocalidades")
            )
        else:
            resumen_export = pd.DataFrame()

        fig_bar_export = None
        if not resumen_export.empty:
            fig_bar_export = px.bar(
                resumen_export,
                x="Provincia",
                y="CantidadLocalidades",
                color="CCTE",
                text="CantidadLocalidades",
                barmode="group",
                title="Localidades por Provincia y CCTE (ámbito del informe)",
                color_discrete_sequence=px.colors.qualitative.Set2
            )
            fig_bar_export.update_layout(template="plotly_white")

        # ========= OPCIONES DE EXPORTACIÓN =========
        col_exp1, _ = st.columns(2)
        formato = col_exp1.radio("Formato de exportación", ["Word (.docx)", "PDF (.pdf)"], horizontal=True)

        if st.button("📄 Generar Informe"):
            localidad_nombre = localidad_seleccionada or "General"
            fecha_str = datetime.now().strftime("%Y%m%d_%H%M")

            # ============================================================
            # 🧾 WORD (sin header azul, con logo y tablas)
            # ============================================================
            if formato == "Word (.docx)":
                doc = Document()

                # Logo arriba del informe (sin header azul)
                if os.path.exists("logo_enacom.png"):
                    doc.add_picture("logo_enacom.png", width=Inches(2.5))

                doc.add_heading(f"Informe de Mediciones RNI - {localidad_nombre}", level=1)
                doc.add_paragraph(f"Ámbito del informe: {titulo_scope}")
                doc.add_paragraph(f"Fecha de generación: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
                doc.add_paragraph(f"Total de puntos medidos: {len(df_export)}")

                if pd.notna(max_resultado):
                    p_max = doc.add_paragraph()
                    p_max.add_run("Resultado máximo registrado: ").bold = True
                    p_max.add_run(f"{max_resultado:.2f} V/m")
                    if max_resultado_pct is not None:
                        p_max.add_run(f" ({max_resultado_pct:.2f} % del límite)")

                    p_ub = doc.add_paragraph()
                    p_ub.add_run("Ubicación del máximo: ").bold = True
                    p_ub.add_run(f"{localidad_max}, {provincia_max} (CCTE {ccte_max})")

                    if fecha_hora_max is not None:
                        p_fm = doc.add_paragraph()
                        p_fm.add_run("Fecha y hora del máximo: ").bold = True
                        p_fm.add_run(str(fecha_hora_max))

                if fecha_min and fecha_max_med:
                    p_f = doc.add_paragraph()
                    p_f.add_run("Rango de fechas de medición: ").bold = True
                    p_f.add_run(f"{fecha_min.strftime('%d/%m/%Y')} a {fecha_max_med.strftime('%d/%m/%Y')}")

                if tiempo_total_trabajado.total_seconds() > 0:
                    p_t = doc.add_paragraph()
                    p_t.add_run("Tiempo total estimado de medición: ").bold = True
                    p_t.add_run(format_timedelta_long(tiempo_total_trabajado))

                if sondas_uniq:
                    p_s = doc.add_paragraph()
                    p_s.add_run("Sondas utilizadas: ").bold = True
                    p_s.add_run(", ".join(sondas_uniq))

                doc.add_paragraph(" ")

                # --- Gráfico principal (ámbito actual) ---
                if fig_bar_export is not None:
                    img_bytes = BytesIO()
                    pio.write_image(fig_bar_export, img_bytes, format="png")
                    img_bytes.seek(0)
                    doc.add_picture(img_bytes, width=Inches(5.5))
                    doc.add_paragraph("Gráfico de Localidades por Provincia y CCTE (ámbito del informe).")

                # --- Desglose por mes (tabla) ---
                if not resumen_mensual_export.empty:
                    doc.add_heading("Desglose por mes", level=2)
                    table = doc.add_table(rows=1, cols=5)
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = "Mes"
                    hdr_cells[1].text = "Puntos"
                    hdr_cells[2].text = "Horas trabajadas"
                    hdr_cells[3].text = "Fecha inicio"
                    hdr_cells[4].text = "Fecha fin"

                    for _, row in resumen_mensual_export.iterrows():
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(row["Mes"])
                        row_cells[1].text = str(row["Cantidad_puntos"])
                        row_cells[2].text = row["Horas_trabajadas"]
                        fi = row["Fecha_inicio"]
                        ff = row["Fecha_fin"]
                        row_cells[3].text = fi.strftime("%d/%m/%Y %H:%M") if pd.notna(fi) else "-"
                        row_cells[4].text = ff.strftime("%d/%m/%Y %H:%M") if pd.notna(ff) else "-"

                # --- Tabla de expedientes ---
                if not expedientes_df.empty:
                    doc.add_heading("Resumen por expediente", level=2)
                    table_e = doc.add_table(rows=1, cols=6)
                    hdr = table_e.rows[0].cells
                    hdr[0].text = "Expediente"
                    hdr[1].text = "Puntos"
                    hdr[2].text = "Max (V/m)"
                    hdr[3].text = "CCTE"
                    hdr[4].text = "Provincias"
                    hdr[5].text = "Localidades"

                    for _, row in expedientes_df.iterrows():
                        r = table_e.add_row().cells
                        r[0].text = str(row["Expediente"])
                        r[1].text = str(row["Cantidad_puntos"])
                        r[2].text = f"{row['Max_Vm']:.2f}" if pd.notna(row["Max_Vm"]) else "-"
                        r[3].text = str(row["CCTE"])
                        r[4].text = str(row["Provincias"])
                        r[5].text = str(row["Localidades"])

                ruta_doc = f"Informe_RNI_{localidad_nombre}_{fecha_str}.docx"
                doc.save(ruta_doc)
                with open(ruta_doc, "rb") as f:
                    st.download_button(
                        label="⬇️ Descargar Informe Word",
                        data=f,
                        file_name=ruta_doc,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

            # ============================================================
            # 📘 PDF (sin header azul, con desglose)
            # ============================================================
            else:
                ruta_pdf = f"Informe_RNI_{localidad_nombre}_{fecha_str}.pdf"
                buffer = BytesIO()
                pdf = SimpleDocTemplate(buffer, pagesize=A4)
                styles = getSampleStyleSheet()
                style_title = styles["Title"]
                style_sub = styles["Heading2"]
                style_normal = styles["Normal"]

                story = []

                # Logo si está disponible
                if os.path.exists("logo_enacom.png"):
                    story.append(RLImage("logo_enacom.png", width=200, height=60))
                    story.append(Spacer(1, 12))

                story.append(Paragraph("Informe de Mediciones RNI", style_title))
                story.append(Spacer(1, 6))
                story.append(Paragraph(f"Ámbito del informe: {titulo_scope}", style_sub))
                story.append(Spacer(1, 12))

                story.append(Paragraph(f"<b>Localidad seleccionada:</b> {localidad_nombre}", style_normal))
                story.append(Paragraph(f"<b>Fecha de generación:</b> {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}", style_normal))
                story.append(Paragraph(f"<b>Total de puntos medidos:</b> {len(df_export)}", style_normal))

                if pd.notna(max_resultado):
                    story.append(Paragraph(
                        f"<b>Resultado máximo registrado:</b> {max_resultado:.2f} V/m"
                        + (f" ({max_resultado_pct:.2f} % del límite)" if max_resultado_pct is not None else ""),
                        style_normal
                    ))
                    story.append(Paragraph(
                        f"<b>Ubicación del máximo:</b> {localidad_max}, {provincia_max} (CCTE {ccte_max})",
                        style_normal
                    ))
                    if fecha_hora_max is not None:
                        story.append(Paragraph(
                            f"<b>Fecha y hora del máximo:</b> {fecha_hora_max}",
                            style_normal
                        ))

                if fecha_min and fecha_max_med:
                    story.append(Paragraph(
                        f"<b>Rango de fechas de medición:</b> {fecha_min.strftime('%d/%m/%Y')} a {fecha_max_med.strftime('%d/%m/%Y')}",
                        style_normal
                    ))

                if tiempo_total_trabajado.total_seconds() > 0:
                    story.append(Paragraph(
                        f"<b>Tiempo total estimado de medición:</b> {format_timedelta_long(tiempo_total_trabajado)}",
                        style_normal
                    ))

                if sondas_uniq:
                    story.append(Paragraph(
                        f"<b>Sondas utilizadas:</b> {', '.join(sondas_uniq)}",
                        style_normal
                    ))

                story.append(Spacer(1, 16))

                # Gráfico (si hay)
                if fig_bar_export is not None:
                    img_bytes = BytesIO()
                    pio.write_image(fig_bar_export, img_bytes, format="png")
                    img_bytes.seek(0)
                    story.append(RLImage(img_bytes, width=400, height=250))
                    story.append(Paragraph("Gráfico de Localidades por Provincia y CCTE (ámbito del informe)", styles["Italic"]))
                    story.append(Spacer(1, 16))

                # Desglose mensual (en texto)
                if not resumen_mensual_export.empty:
                    story.append(Paragraph("<b>Desglose por mes</b>", style_sub))
                    story.append(Spacer(1, 6))
                    for _, row in resumen_mensual_export.iterrows():
                        fi = row["Fecha_inicio"]
                        ff = row["Fecha_fin"]
                        texto = (
                            f"Mes {row['Mes']}: {row['Cantidad_puntos']} puntos, "
                            f"horas trabajadas: {row['Horas_trabajadas']}, "
                            f"localidades: {row['Localidades_trabajadas']}. "
                        )
                        if pd.notna(fi) and pd.notna(ff):
                            texto += f"({fi.strftime('%d/%m/%Y %H:%M')} a {ff.strftime('%d/%m/%Y %H:%M')})"
                        story.append(Paragraph(texto, style_normal))
                    story.append(Spacer(1, 12))

                # Tabla de expedientes (en texto)
                if not expedientes_df.empty:
                    story.append(Paragraph("<b>Resumen por expediente</b>", style_sub))
                    story.append(Spacer(1, 6))
                    for _, row in expedientes_df.iterrows():
                        texto = (
                            f"Expediente {row['Expediente']}: "
                            f"{row['Cantidad_puntos']} puntos, "
                            f"máx {row['Max_Vm']:.2f} V/m, "
                            f"CCTE: {row['CCTE']}, "
                            f"Provincias: {row['Provincias']}, "
                            f"Localidades: {row['Localidades']}."
                        )
                        story.append(Paragraph(texto, style_normal))
                    story.append(Spacer(1, 12))

                pdf.build(story)
                buffer.seek(0)
                st.download_button(
                    label="⬇️ Descargar Informe PDF",
                    data=buffer,
                    file_name=ruta_pdf,
                    mime="application/pdf"
                )

# ============================================================
# 📊 TABLA MAESTRA
# ============================================================

st.header("📊 Tabla Maestra de Mediciones RNI")

if st.session_state["tabla_maestra"].empty:
    st.info("La tabla maestra está vacía. Cargá archivos a la izquierda.")
else:
    total_registros = len(st.session_state["tabla_maestra"])
    st.caption(f"🗂️ Registros totales: **{total_registros:,}**")
    with st.expander("📂 Mostrar / Ocultar tabla maestra", expanded=False):
        df_maestra = st.session_state["tabla_maestra"].copy().dropna(axis=1, how="all")
        st.dataframe(df_maestra.reset_index(drop=True), width="stretch")
        if st.button("💾 Exportar tabla a Excel"):
            df_maestra.to_excel("tabla_maestra.xlsx", index=False)
            with open("tabla_maestra.xlsx", "rb") as f:
                st.download_button("⬇️ Descargar Excel", data=f, file_name="tabla_maestra.xlsx")
